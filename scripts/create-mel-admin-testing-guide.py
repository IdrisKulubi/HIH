from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "src" / "docs" / "BIRE_MEL_Admin_Testing_and_Acceptance_Guide.docx"

BLUE = "009FE3"
DARK_BLUE = "14547A"
INK = "1F2937"
MUTED = "5B6573"
PALE_BLUE = "E8F5FB"
PALE_GRAY = "F3F5F7"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF7E5"
WHITE = "FFFFFF"
BORDER = "C9D5DF"
RED = "A52525"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_font(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_text(cell, text, bold=False, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_font(run, size=size, color=color, bold=bold)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), BORDER)
        borders.append(border)
    p_pr.append(borders)
    r1 = p.add_run(f"{label}: ")
    set_font(r1, size=10, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=10, color=INK)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Field", bold=True, color=WHITE, size=9.2)
    set_cell_text(table.rows[0].cells[1], "Details", bold=True, color=WHITE, size=9.2)
    set_cell_shading(table.rows[0].cells[0], DARK_BLUE)
    set_cell_shading(table.rows[0].cells[1], DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=DARK_BLUE, size=9.5)
        set_cell_text(cells[1], value, size=9.5)
        set_cell_shading(cells[0], PALE_BLUE)
    set_table_geometry(table, [2500, 6860])
    return table


def add_test_table(doc, rows):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["ID", "Manager action", "Expected result", "Evidence", "Result"]
    widths = [650, 3210, 2900, 1600, 1000]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE, size=8.8,
                      align=WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(table.rows[0].cells[idx], DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for test_id, action, expected, evidence in rows:
        cells = table.add_row().cells
        values = [test_id, action, expected, evidence, "Pass / Fail / N/A"]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=8.6, align=align)
            if idx == 0:
                set_cell_shading(cells[idx], PALE_GRAY)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BIRE MEL | ADMIN TESTING AND ACCEPTANCE GUIDE")
    set_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BIRE Programme  |  Page ")
    set_font(run, size=9, color=MUTED)
    add_field(p, "PAGE")


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("BIRE PROGRAMME")
    set_font(r, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("MEL System\nAdmin Testing and Acceptance Guide")
    set_font(r, size=27, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("A practical end-to-end guide for management review, usability testing, and operational sign-off")
    set_font(r, size=13, color=MUTED)

    add_key_value_table(doc, [
        ("Primary user", "Programme manager or designated tester with the admin role"),
        ("Coverage", "MEL Phases 1-5, security, usability, reporting, operations, and rollout"),
        ("Document date", "31 July 2026"),
        ("Recommended environment", "Isolated test or staging environment with representative data"),
        ("System URL", "____________________________________________"),
        ("Test lead", "____________________________________________"),
    ])
    add_body(doc, "Use this document as the official UAT record. Enter Pass, Fail, or N/A in each result cell and attach the named evidence. Do not use confidential production data for testing.", after=0)
    add_page_break(doc)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. How the manager should use this guide", 1)
    add_body(doc, "Complete the tests in order. Phase 1 provides the configuration used by later phases; Phase 2 creates reports; Phase 3 approves trusted data; Phase 4 calculates and presents results; Phase 5 controls instruments, imports, and rollout.")
    add_callout(doc, "Important", "An admin account can manage all MEL workspaces, but separation-of-duties tests require at least one additional account. A user must not approve a submission they created or promoted.", PALE_AMBER)
    add_key_value_table(doc, [
        ("Pass", "The expected result occurs, data is correct, and required evidence is captured."),
        ("Fail", "The result is incorrect, blocked unexpectedly, insecure, misleading, or difficult to use."),
        ("N/A", "The test is outside the agreed deployment scope; record the reason."),
        ("Stop condition", "Stop testing and escalate immediately for data loss, unauthorized access, exposed secrets, or incorrect official totals."),
    ])

    add_heading(doc, "2. Access and first-time setup", 1)
    add_test_table(doc, [
        ("A01", "Open /login and sign in with an active admin account.", "The admin dashboard opens and shows the Management panel.", "Dashboard screenshot"),
        ("A02", "Confirm Management shows MEL Configuration, Quarterly Monitoring, MEL Review Queue, MEL Reporting, MEL Instruments, MEL Imports, and MEL Operations.", "Every named link is visible and opens without a permissions error.", "Navigation screenshots"),
        ("A03", "Open /admin/mel/operations. Keep the stage as internal test; enable Collection, Imports, and Reporting; save.", "The switches remain enabled after refresh and an audit/operational record is created.", "Rollout-control screenshot"),
        ("A04", "Sign out, then try a user without MEL access on an MEL URL.", "The user is redirected or receives a clear access-denied response; no MEL data is exposed.", "Access-denied screenshot"),
    ])

    add_heading(doc, "3. Test data to prepare", 1)
    add_key_value_table(doc, [
        ("Users", "Admin manager, EDO/data collector, REDO reviewer, MEL approver, and one unauthorized user."),
        ("Enterprises", "At least three enterprises across two counties, two sectors, and more than one programme track."),
        ("Periods", "One open quarter, one closed quarter, and a previous quarter with approved results."),
        ("Indicators", "Representative Impact, LT1-LT4, Output 1-4, percentage, count, currency, and lower-is-better indicators."),
        ("Evidence", "A valid PDF/image, an unsupported file type, and an oversized file for negative testing."),
        ("Import payloads", "One valid record, one duplicate, one unknown enterprise, one invalid period, and one record with missing required values."),
    ])
    add_callout(doc, "Data protection", "Use synthetic names, contacts, coordinates, documents, and financial values. Never paste a webhook secret into screenshots, email, chat, or the defect log.")

    add_page_break(doc)
    add_heading(doc, "4. Phase 1 - MEL foundation and reporting controls", 1)
    add_body(doc, "Purpose: confirm that the admin can control reporting periods, indicators, targets, inclusion data, and audited changes from MEL Configuration.")
    add_test_table(doc, [
        ("P1-01", "Open /admin/mel and inspect the reporting-period list.", "Period codes, labels, dates, sequence, and status are clear and correctly ordered.", "Period-list screenshot"),
        ("P1-02", "Create or update a test reporting period with valid dates.", "The period saves, displays correctly after refresh, and records an audit event.", "Before/after screenshot"),
        ("P1-03", "Attempt invalid dates or a duplicate period code.", "The system rejects the change with a clear, actionable message.", "Validation screenshot"),
        ("P1-04", "Inspect indicator definitions, disaggregation settings, calculation rules, and targets.", "Definitions match the approved ITT and the interface distinguishes editable configuration from calculated values.", "Indicator sample"),
        ("P1-05", "Close a test reporting period and attempt to start or edit collection in it.", "Collection is blocked while historical records remain readable.", "Blocked-action screenshot"),
        ("P1-06", "Check target totals and disaggregated target distribution.", "Disaggregated targets reconcile with the approved overall target or show a clear exception.", "Reconciliation note"),
    ])
    add_callout(doc, "Phase 1 acceptance", "Pass when configuration is accurate, invalid states are rejected, closed periods are protected, and sensitive changes are auditable.", PALE_GREEN)

    add_heading(doc, "5. Phase 2 - Quarterly enterprise monitoring", 1)
    add_body(doc, "Purpose: confirm that an enterprise report can be started, saved safely, validated, supported by evidence, and submitted exactly once for the selected period.")
    add_test_table(doc, [
        ("P2-01", "Open /admin/mel/monitoring; filter/search and select an enterprise plus the open period.", "The correct enterprise context opens and an existing submission is reused instead of duplicated.", "Queue and form screenshot"),
        ("P2-02", "Complete sections A-I with representative capacity, profitability, jobs, markets, finance, green growth, partnership, and feedback values.", "Fields accept the correct data types; derived totals update consistently; help and labels are understandable.", "Completed draft"),
        ("P2-03", "Save a partial draft, navigate away, then reopen it.", "All saved values return unchanged and the record remains a draft.", "Draft restoration screenshot"),
        ("P2-04", "Enter negative counts, inconsistent job totals, missing mandatory values, and an invalid percentage.", "Submission is blocked and each problem is identified near the affected field or in a clear summary.", "Validation evidence"),
        ("P2-05", "Upload acceptable evidence, then test an unsupported/oversized file.", "Valid evidence is linked to the submission; invalid files are rejected safely.", "Evidence list and error"),
        ("P2-06", "Submit a valid report and try to create another for the same enterprise and period.", "The status advances to submitted and the duplicate is prevented.", "Status screenshot"),
        ("P2-07", "Review the form at desktop and narrow/mobile widths using keyboard navigation.", "Content remains readable, focus is visible, controls are usable, and no content is clipped.", "Responsive screenshots"),
    ])
    add_callout(doc, "Phase 2 acceptance", "Pass when drafts are reliable, calculations and validation are correct, evidence is protected, duplicates are blocked, and submission status is clear.", PALE_GREEN)

    add_page_break(doc)
    add_heading(doc, "6. Phase 3 - Review, DQA, evidence, and learning", 1)
    add_body(doc, "Purpose: prove that only reviewed and approved data becomes trusted, while corrections, evidence decisions, DQA findings, notifications, and learning actions remain traceable.")
    add_test_table(doc, [
        ("P3-01", "Open /admin/mel/review and inspect queue filters, status, age, origin, enterprise, and period.", "The queue is understandable and the test submission is visible to the appropriate reviewer.", "Queue screenshot"),
        ("P3-02", "Open a submission; compare its answers, calculations, history, and evidence with the original report.", "The review detail is complete and no source data is silently changed.", "Review-detail screenshot"),
        ("P3-03", "As REDO/admin, return an EDO-originated report with a correction reason.", "A meaningful reason is required; status changes correctly; the collector receives a notification.", "Status/history evidence"),
        ("P3-04", "Correct and resubmit the returned report, then complete REDO validation and MEL approval using separate users.", "The full status sequence is recorded and self-approval is blocked.", "Audit timeline"),
        ("P3-05", "Record completeness, consistency, plausibility, and timeliness DQA results, including a failed check.", "Scores/findings are saved; failed checks remain visible and actionable.", "DQA screenshot"),
        ("P3-06", "Verify and reject different evidence items with reasons.", "Each evidence decision is independent, auditable, and visible in /admin/mel/evidence.", "Evidence decisions"),
        ("P3-07", "Create, assign, update, and close a learning action in /admin/mel/learning.", "Owner, due date, priority, history, and closure evidence are retained.", "Learning-action record"),
        ("P3-08", "Attempt approval from the same account that created/promoted the submission.", "Approval is refused with a clear separation-of-duties message.", "Denied-action screenshot"),
    ])
    add_callout(doc, "Phase 3 acceptance", "Pass when workflow states cannot be skipped, returns require reasons, self-approval is prevented, DQA/evidence are auditable, and notifications reach the right role.", PALE_GREEN)

    add_page_break(doc)
    add_heading(doc, "7. Phase 4 - ITT, dashboards, GIS, and reporting", 1)
    add_body(doc, "Purpose: confirm that official results use approved data only, match manual calculations, disclose lineage, apply filters consistently, and export safely.")
    add_test_table(doc, [
        ("P4-01", "Open /admin/mel/programme-results; create a programme-level entry, attach its evidence URL where required, and approve it.", "Draft values are excluded from official totals; approved values become eligible for calculation.", "Entry and status evidence"),
        ("P4-02", "Open /admin/mel/reporting; select a period and recalculate.", "Actuals, targets, achievement, traffic lights, calculation version, source count, and timestamp appear.", "Dashboard screenshot"),
        ("P4-03", "Manually calculate representative Impact, LT, Output, percentage, count, and currency indicators.", "System actuals match the signed manual reconciliation under identical filters.", "Reconciliation worksheet"),
        ("P4-04", "Open an indicator source-count/lineage link.", "The manager can trace the result to approved source records and see exclusions/missing-data notes.", "Lineage screenshot"),
        ("P4-05", "Apply track, county, sector, youth, gender, disability, refugee, and enterprise filters where data exists.", "Every dashboard element and export uses the same filter context without mixing populations.", "Filtered results"),
        ("P4-06", "Open /admin/mel/reporting/data-quality.", "Missing, stale, invalid, or unreconciled data is clearly distinguished from zero performance.", "Data-quality screenshot"),
        ("P4-07", "Open /admin/mel/gis and inspect verified, invalid, clustered, and rounded coordinates.", "Only protected/verified locations are mapped; precise coordinates are not exposed to unauthorized users.", "GIS screenshot"),
        ("P4-08", "Export identical filters to CSV and Excel.", "Dashboard, CSV, and Excel totals agree; metadata identifies period, filters, timestamp, and trusted-data rule.", "Export files"),
        ("P4-09", "Return, reopen, or void an approved source; recalculate.", "The source leaves official totals until it is approved again, with no change to unrelated results.", "Before/after totals"),
    ])
    add_callout(doc, "Phase 4 acceptance", "Pass only after representative ITT values reconcile exactly, trusted-data rules hold, filters/exports agree, GIS privacy is preserved, and lineage is understandable.", PALE_GREEN)

    add_page_break(doc)
    add_heading(doc, "8. Phase 5 - Configurable tools, imports, and rollout", 1)
    add_body(doc, "Purpose: validate safe instrument versioning, controlled integration, duplicate/quarantine handling, operational visibility, and evidence-based rollout progression.")
    add_heading(doc, "8.1 Configurable instruments", 2)
    add_test_table(doc, [
        ("P5-01", "Open /admin/mel/instruments and create TEST_BASELINE as a draft.", "Version 1 is created as an editable draft with a unique code.", "Instrument screenshot"),
        ("P5-02", "Add sections and questions covering text, number, currency, percentage, date, boolean, select, file, required, evidence, and indicator mapping.", "Question settings persist and incompatible mappings/options are rejected.", "Question-definition evidence"),
        ("P5-03", "Create a conditional question that depends on another answer; test a missing dependency and a circular rule.", "Valid visibility works; missing or circular dependencies prevent publishing.", "Validation messages"),
        ("P5-04", "Validate and publish a complete instrument, then attempt to edit it.", "The published version is immutable and records publisher/time/effective period.", "Published version screenshot"),
        ("P5-05", "Create the next version and retire a published version with a meaningful reason.", "The new draft is independently editable; retirement requires and records a reason.", "Version history"),
    ])

    add_heading(doc, "8.2 Imports and quarantine", 2)
    add_test_table(doc, [
        ("P5-06", "Open /admin/mel/imports; create a connection to a published instrument and copy the generated secret once.", "The secret is shown once only; reopening the page never reveals it.", "Redacted connection screenshot"),
        ("P5-07", "Create a field mapping for enterprise, reporting period, external submission ID, and approved question/indicator fields.", "A new mapping version becomes active and the previous version remains traceable.", "Mapping summary"),
        ("P5-08", "Run the controlled test payload with a valid record.", "The record validates, normalized data excludes unmapped/private fields, and an idempotency key is recorded.", "Validated record"),
        ("P5-09", "Send the same external submission twice.", "The duplicate is detected and does not create a second submission.", "Duplicate response/event"),
        ("P5-10", "Test unknown enterprise, invalid period, missing fields, oversized payload, and wrong secret.", "Invalid records are quarantined or rejected with safe messages; no secret/raw sensitive data appears in events.", "Quarantine/event evidence"),
        ("P5-11", "Promote a validated record to review, then inspect the review queue.", "Promotion creates a submitted record but never auto-approves it; normal review controls apply.", "Queue and audit evidence"),
    ])

    add_heading(doc, "8.3 Operations and rollout", 2)
    add_test_table(doc, [
        ("P5-12", "Open /admin/mel/operations and inspect operational events and correlation IDs.", "Events are useful for investigation without exposing answers, files, tokens, or secrets.", "Operations screenshot"),
        ("P5-13", "Disable each feature flag and try collection, imports, and reporting; then re-enable it.", "The corresponding feature is blocked safely when disabled and works when enabled.", "Flag test evidence"),
        ("P5-14", "Attempt to move from internal test to pilot before completing required gates.", "Advancement is blocked and identifies ITT reconciliation, security review, and internal UAT requirements.", "Blocked-stage screenshot"),
        ("P5-15", "Record owner and evidence for a gate, then mark it passed or waived.", "Evidence is mandatory, the change is auditable, and eligible stage progression becomes available.", "Gate record"),
        ("P5-16", "Move a stage backwards without a reason, then retry with a meaningful rollback reason.", "The first attempt is blocked; the second succeeds and records the reason.", "Rollback evidence"),
    ])
    add_callout(doc, "Phase 5 acceptance", "Pass when published instruments are immutable, imports are authenticated/idempotent/quarantined, promoted data follows review, secrets remain protected, and rollout gates cannot be bypassed.", PALE_GREEN)

    add_page_break(doc)
    add_heading(doc, "9. Cross-cutting usability and security review", 1)
    add_test_table(doc, [
        ("X01", "Complete the critical journeys without technical assistance and note unclear labels or unnecessary steps.", "The manager can understand current status, next action, and consequences on every screen.", "Usability notes"),
        ("X02", "Use keyboard-only navigation through forms, tables, dialogs, filters, and buttons.", "Focus order is logical, visible, and no control requires a mouse.", "Accessibility notes"),
        ("X03", "Test at approximately 375 px, 768 px, and desktop width.", "No horizontal clipping, inaccessible action, unreadable table, or overlapping content appears.", "Responsive screenshots"),
        ("X04", "Try direct URLs and server actions using view-only or unauthorized accounts.", "Server-side authorization blocks changes even if a URL or request is known.", "Authorization evidence"),
        ("X05", "Review errors generated during invalid inputs, imports, and exports.", "Messages help the user recover but do not disclose stack traces, SQL, tokens, or confidential payloads.", "Error samples"),
        ("X06", "Run agreed expected-volume monitoring, reporting, export, and import tests.", "Response times remain acceptable and failures create actionable operational events.", "Performance results"),
        ("X07", "Perform and document a non-production backup restoration.", "The restored system contains the expected MEL data and passes agreed integrity checks.", "Restore report"),
    ])

    add_heading(doc, "10. Manager acceptance summary", 1)
    summary = doc.add_table(rows=1, cols=5)
    summary.style = "Table Grid"
    headers = ["Area", "Passed", "Failed", "N/A", "Manager notes"]
    for i, value in enumerate(headers):
        set_cell_text(summary.rows[0].cells[i], value, bold=True, color=WHITE, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if i != 4 else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(summary.rows[0].cells[i], DARK_BLUE)
    set_repeat_table_header(summary.rows[0])
    for area in ["Access/setup", "Phase 1", "Phase 2", "Phase 3", "Phase 4", "Phase 5", "Usability/security", "Backup/operations"]:
        cells = summary.add_row().cells
        for i, value in enumerate([area, "", "", "", ""]):
            set_cell_text(cells[i], value, size=9, align=WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(summary, [1800, 850, 850, 850, 5010])

    add_heading(doc, "11. Defect log", 1)
    add_body(doc, "Record every failed test. Critical and high-severity defects must be closed or formally accepted before production approval.")
    defects = doc.add_table(rows=1, cols=6)
    defects.style = "Table Grid"
    headers = ["Defect ID", "Test ID", "Severity", "Description", "Owner", "Status"]
    for i, value in enumerate(headers):
        set_cell_text(defects.rows[0].cells[i], value, bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(defects.rows[0].cells[i], DARK_BLUE)
    set_repeat_table_header(defects.rows[0])
    for _ in range(8):
        cells = defects.add_row().cells
        for cell in cells:
            set_cell_text(cell, "", size=8.5)
            cell.paragraphs[0].add_run("\n")
    set_table_geometry(defects, [1100, 900, 1050, 3410, 1450, 1450])

    add_page_break(doc)
    add_heading(doc, "12. Final sign-off", 1)
    add_callout(doc, "Decision rule", "Approve production use only when critical/high defects are closed, ITT reconciliation is signed, privacy/security checks pass, restore evidence exists, role training is complete, and rollout gates contain valid evidence.", PALE_AMBER)
    add_key_value_table(doc, [
        ("Overall decision", "Approved / Approved with conditions / Not approved"),
        ("Conditions or exclusions", "________________________________________________________________________________\n________________________________________________________________________________"),
        ("Manager name", "____________________________________________"),
        ("Signature", "____________________________________________"),
        ("Date", "____________________________________________"),
        ("MEL lead", "____________________________________________"),
        ("Technical lead", "____________________________________________"),
    ])
    add_body(doc, "Retain this signed guide together with reconciliation workbooks, screenshots, export samples, restore evidence, training attendance, and the final defect register as the MEL acceptance pack.")

    doc.core_properties.title = "BIRE MEL System Admin Testing and Acceptance Guide"
    doc.core_properties.subject = "End-to-end manager UAT and acceptance guide for MEL Phases 1-5"
    doc.core_properties.author = "BIRE Programme"
    doc.core_properties.keywords = "BIRE, MEL, UAT, admin, testing, acceptance"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
